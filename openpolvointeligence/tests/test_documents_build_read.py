"""Build (python-docx) e leitura/digest determinísticos de documentos Word."""

from __future__ import annotations

import base64
import importlib.util

import pytest

from openpolvointeligence.graphs.documents_full.documents_build_logic import (
    DocxBuildError,
    apply_edit_ops,
    build_docx_bytes,
    build_docx_metadata,
    ensure_docx_filename,
)
from openpolvointeligence.graphs.documents_full.documents_read_logic import (
    decode_attachment,
    digest_to_markdown,
    read_document_digest,
)
from openpolvointeligence.graphs.documents_full.documents_spec_logic import normalize_edit_plan

_HAS_DOCX = importlib.util.find_spec("docx") is not None

_SPEC = {
    "filename": "relatorio.docx",
    "title": "Relatório Anual",
    "blocks": [
        {"type": "heading", "level": 1, "text": "Introdução"},
        {"type": "paragraph", "text": "Texto de abertura.", "bold_phrases": ["abertura"]},
        {"type": "bullet_list", "items": ["Ponto A", "Ponto B"]},
        {"type": "table", "headers": ["Métrica", "Valor"], "rows": [["Vendas", "100"]]},
    ],
}


def test_decode_attachment_roundtrip() -> None:
    data = base64.b64encode(b"hello").decode()
    assert decode_attachment({"data_base64": data}) == b"hello"


def test_digest_to_markdown_includes_stats() -> None:
    md = digest_to_markdown(
        {
            "filename": "t.docx",
            "kind": "docx",
            "headings": 2,
            "paragraphs": 5,
            "tables": 1,
            "markdown": "# Título\n\nCorpo.",
            "error": "",
        }
    )
    assert "2 secções" in md
    assert "# Título" in md


def test_ensure_docx_filename() -> None:
    assert ensure_docx_filename("rel") == "rel.docx"
    assert ensure_docx_filename("rel.doc") == "rel.docx"
    assert ensure_docx_filename("") == "documento.docx"


def test_build_docx_metadata() -> None:
    meta = build_docx_metadata(b"bytes", filename="x.docx")
    assert meta["document_kind"] == "docx_result"
    assert base64.b64decode(meta["docx_document_base64"]) == b"bytes"


@pytest.mark.skipif(not _HAS_DOCX, reason="requer python-docx instalado")
def test_build_docx_bytes_and_read_back() -> None:
    import io

    from docx import Document  # type: ignore[import-untyped]

    data = build_docx_bytes(_SPEC)
    assert data[:2] == b"PK"
    doc = Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "Relatório Anual" in texts
    assert "Introdução" in texts
    assert "Texto de abertura." in texts
    assert len(doc.tables) == 1


@pytest.mark.skipif(not _HAS_DOCX, reason="requer python-docx instalado")
def test_read_docx_digest() -> None:
    data = build_docx_bytes(_SPEC)
    digest = read_document_digest(data, filename="relatorio.docx")
    assert digest["kind"] == "docx"
    assert digest["paragraphs"] >= 1
    assert "Introdução" in (digest.get("markdown") or "")


@pytest.mark.skipif(not _HAS_DOCX, reason="requer python-docx instalado")
def test_apply_edit_ops_replace_text() -> None:
    import io

    from docx import Document  # type: ignore[import-untyped]

    data = build_docx_bytes(_SPEC)
    plan = normalize_edit_plan(
        {"ops": [{"op": "replace_text", "find": "abertura", "replace": "conclusão"}]}
    )
    edited = apply_edit_ops(data, plan)
    doc = Document(io.BytesIO(edited))
    joined = " ".join(p.text for p in doc.paragraphs)
    assert "conclusão" in joined


def test_build_docx_without_library() -> None:
    if _HAS_DOCX:
        pytest.skip("python-docx instalado")
    with pytest.raises(DocxBuildError):
        build_docx_bytes(_SPEC)
