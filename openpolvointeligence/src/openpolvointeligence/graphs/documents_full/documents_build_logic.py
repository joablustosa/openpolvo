"""Construção/edição determinística de ficheiros .docx via python-docx (zero alucinação)."""

from __future__ import annotations

import base64
import io
import logging
import re
from typing import Any

_logger = logging.getLogger(__name__)

_MAX_DOCX_BYTES = 8 * 1024 * 1024
_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class DocxBuildError(RuntimeError):
    """python-docx ausente ou falha irrecuperável ao construir o documento."""


def _require_docx() -> Any:
    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.shared import Inches  # type: ignore[import-untyped]
    except ImportError as exc:  # pragma: no cover
        raise DocxBuildError("python-docx não instalado — necessário para gerar .docx.") from exc
    return Document, Inches


def _heading_style(level: int) -> str:
    return "Title" if level <= 0 else f"Heading {min(max(level, 1), 3)}"


def _apply_runs(
    paragraph: Any, text: str, *, bold_phrases: list[str], italic_phrases: list[str]
) -> None:
    if not bold_phrases and not italic_phrases:
        paragraph.add_run(text)
        return
    patterns: list[tuple[str, str]] = []
    for phrase in bold_phrases:
        if phrase:
            patterns.append((re.escape(phrase), "bold"))
    for phrase in italic_phrases:
        if phrase:
            patterns.append((re.escape(phrase), "italic"))
    if not patterns:
        paragraph.add_run(text)
        return
    combined = "|".join(f"({p})" for p, _ in patterns)
    parts = re.split(combined, text)
    for part in parts:
        if part is None or part == "":
            continue
        run = paragraph.add_run(part)
        if part in bold_phrases:
            run.bold = True
        if part in italic_phrases:
            run.italic = True


def _add_block(doc: Any, block: dict[str, Any]) -> None:
    kind = block.get("type")
    if kind == "heading":
        level = int(block.get("level") or 1)
        doc.add_paragraph(str(block.get("text") or ""), style=_heading_style(level))
    elif kind == "paragraph":
        p = doc.add_paragraph(style="Normal")
        _apply_runs(
            p,
            str(block.get("text") or ""),
            bold_phrases=list(block.get("bold_phrases") or []),
            italic_phrases=list(block.get("italic_phrases") or []),
        )
    elif kind == "bullet_list":
        for item in block.get("items") or []:
            doc.add_paragraph(str(item), style="List Bullet")
    elif kind == "numbered_list":
        for item in block.get("items") or []:
            doc.add_paragraph(str(item), style="List Number")
    elif kind == "table":
        headers = list(block.get("headers") or [])
        rows = list(block.get("rows") or [])
        ncols = max(len(headers), max((len(r) for r in rows), default=0), 1)
        nrows = (1 if headers else 0) + len(rows)
        table = doc.add_table(rows=nrows, cols=ncols)
        table.style = "Table Grid"
        r_off = 0
        if headers:
            for c, h in enumerate(headers):
                cell = table.rows[0].cells[c]
                cell.text = str(h)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            r_off = 1
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                if c_idx < ncols:
                    table.rows[r_idx + r_off].cells[c_idx].text = str(val)


def _apply_page_setup(doc: Any, page_setup: dict[str, float]) -> None:
    _, Inches = _require_docx()
    section = doc.sections[0]
    mapping = {
        "margin_top_in": "top_margin",
        "margin_bottom_in": "bottom_margin",
        "margin_left_in": "left_margin",
        "margin_right_in": "right_margin",
    }
    for key, attr in mapping.items():
        val = page_setup.get(key)
        if val is not None:
            try:
                setattr(section, attr, Inches(float(val)))
            except (TypeError, ValueError):
                continue
    # Margens padrão profissionais (1") se nenhuma especificada.
    if not page_setup:
        for attr in mapping.values():
            setattr(section, attr, Inches(1.0))


def build_docx_bytes(spec: dict[str, Any]) -> bytes:
    Document, _ = _require_docx()
    doc = Document()
    title = str(spec.get("title") or "").strip()
    if title:
        doc.add_paragraph(title, style="Title")
    blocks = spec.get("blocks") or []
    if not blocks:
        doc.add_paragraph("(documento vazio)", style="Normal")
    for block in blocks:
        _add_block(doc, block)
    _apply_page_setup(doc, spec.get("page_setup") or {})
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def load_document_from_bytes(data: bytes) -> Any:
    Document, _ = _require_docx()
    return Document(io.BytesIO(data))


def _append_block_to_doc(doc: Any, block: dict[str, Any]) -> None:
    _add_block(doc, block)


def apply_edit_ops(data: bytes, plan: dict[str, Any]) -> bytes:
    """Carrega o documento e aplica operações mínimas, preservando o restante."""
    doc = load_document_from_bytes(data)
    paragraphs = doc.paragraphs
    tables = doc.tables

    for op in plan.get("ops") or []:
        kind = op.get("op")
        if kind == "replace_text":
            find = str(op.get("find") or "")
            replace = str(op.get("replace") or "")
            if not find:
                continue
            for p in paragraphs:
                if find in p.text:
                    p.text = p.text.replace(find, replace)
        elif kind == "insert_paragraph":
            idx = int(op.get("paragraph_index") or len(paragraphs))
            text = str(op.get("text") or "")
            if 0 <= idx < len(paragraphs):
                paragraphs[idx].insert_paragraph_before(text)
            else:
                doc.add_paragraph(text, style="Normal")
        elif kind == "insert_heading":
            level = int(op.get("level") or 1)
            text = str(op.get("text") or "")
            doc.add_paragraph(text, style=_heading_style(level))
        elif kind == "append_block":
            block = op.get("block")
            if isinstance(block, dict):
                _append_block_to_doc(doc, block)
        elif kind == "delete_paragraph":
            try:
                idx = int(op.get("paragraph_index"))
                if 0 <= idx < len(paragraphs):
                    p = paragraphs[idx]._element  # noqa: SLF001
                    p.getparent().remove(p)
            except (TypeError, ValueError, IndexError):
                continue
        elif kind == "set_table_cell":
            try:
                t_idx = int(op.get("table_index") or 0)
                row = int(op.get("row") or 0)
                col = int(op.get("col") or 0)
                value = str(op.get("value") or "")
                if 0 <= t_idx < len(tables):
                    tables[t_idx].rows[row].cells[col].text = value
            except (TypeError, ValueError, IndexError):
                continue
        elif kind == "add_table_row":
            try:
                t_idx = int(op.get("table_index") or 0)
                values = [str(v) for v in (op.get("values") or [])]
                if 0 <= t_idx < len(tables):
                    row = tables[t_idx].add_row()
                    for c, val in enumerate(values):
                        if c < len(row.cells):
                            row.cells[c].text = val
            except (TypeError, ValueError, IndexError):
                continue
        elif kind == "insert_table":
            block = op.get("block") if isinstance(op.get("block"), dict) else op
            if isinstance(block, dict):
                tbl = block if block.get("type") == "table" else {"type": "table", **block}
                _append_block_to_doc(doc, tbl)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_docx_metadata(
    docx_bytes: bytes,
    *,
    filename: str,
    extra: dict[str, Any] | None = None,
) -> dict[str, Any]:
    meta: dict[str, Any] = {
        "document_kind": "docx_result",
        "document_format": _DOCX_MIME,
        "docx_export_suggested_filename": filename,
    }
    if extra:
        meta.update(extra)
    if len(docx_bytes) <= _MAX_DOCX_BYTES:
        meta["docx_document_base64"] = base64.b64encode(docx_bytes).decode("ascii")
        meta["docx_size_bytes"] = len(docx_bytes)
    else:
        meta["docx_too_large"] = True
        meta["docx_size_bytes"] = len(docx_bytes)
    return meta


def ensure_docx_filename(name: str, *, default: str = "documento.docx") -> str:
    n = (name or "").strip() or default
    if n.lower().endswith(".doc"):
        n = n[:-4]
    if not n.lower().endswith(".docx"):
        n = n.rsplit(".", 1)[0] if "." in n else n
        n = f"{n}.docx"
    return n
